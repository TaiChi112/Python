docWrite(doc,table_data2,h2,head)
doc.save('./Assignment2/ThaiPBS_Webscrapping.docx')

# --- Imports ---
import requests
from bs4 import BeautifulSoup
import lxml
from docx import Document
import os

# --- Constants ---
WIKI_URL = 'https://th.wikipedia.org/wiki/%E0%B8%AA%E0%B8%96%E0%B8%B2%E0%B8%99%E0%B8%B5%E0%B9%82%E0%B8%97%E0%B8%A3%E0%B8%97%E0%B8%B1%E0%B8%A8%E0%B8%99%E0%B9%8C%E0%B9%84%E0%B8%97%E0%B8%A2%E0%B8%9E%E0%B8%B5%E0%B8%9A%E0%B8%B5%E0%B9%80%E0%B8%AD%E0%B8%AA'
OUTPUT_DOC_PATH = './Assignment2/ThaiPBS_Webscrapping.docx'

# --- Web Scraping ---
response = requests.get(WIKI_URL)
soup = BeautifulSoup(response.text, 'html.parser')

# --- Extract Director Table ---
director_table = soup.find('table', {'class': 'toccolours'})
director_heading = soup.find('h3', {'id': 'ผู้อำนวยการสถานีฯ'})
director_columns = director_table.find_all('th')
director_rows = director_table.find_all('tr')[2:]  # Skip header rows

# --- Extract News Schedule Table ---
news_table = soup.find('table', {'class': 'wikitable'})
news_heading = soup.find('h2', {'id': 'รายการข่าวของสถานี_ฯ'})
news_columns = news_table.find_all('th')
news_rows = news_table.find_all('tr')[1:]  # Skip header row

# --- Document Creation ---
doc = Document()

def add_table_to_doc(document, row_data, heading_tag, column_tags):
    """
    Add a table to the Word document with a heading.
    :param document: python-docx Document object
    :param row_data: list of BeautifulSoup <tr> tags (table rows)
    :param heading_tag: BeautifulSoup tag for the table heading
    :param column_tags: list of BeautifulSoup <th> tags (column headers)
    """
    document.add_heading(heading_tag.get_text(strip=True), level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for idx, col_tag in enumerate(column_tags):
        header_cells[idx].text = col_tag.get_text(strip=True)

    for row_tag in row_data:
        # Remove superscript references
        for sup in row_tag.find_all('sup'):
            sup.decompose()
        cells = row_tag.find_all('td')
        row_cells = table.add_row().cells
        if len(cells) == 1:
            row_cells[0].text = ""
            row_cells[1].text = cells[0].get_text(strip=True)
        elif len(cells) == 2:
            row_cells[0].text = cells[0].get_text(strip=True)
            # Handle multi-line content in the second cell
            text_lines = [text.strip() for text in cells[1].stripped_strings]
            if len(text_lines) == 1:
                row_cells[1].text = cells[1].get_text(strip=True)
            else:
                row_cells[1].text = " ".join(text_lines)

# --- Add Tables to Document ---
add_table_to_doc(doc, director_rows, director_heading, director_columns)
add_table_to_doc(doc, news_rows, news_heading, news_columns)

# --- Save Document ---
doc.save(OUTPUT_DOC_PATH)





